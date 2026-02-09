from docx import Document
import re
from datetime import datetime

# ==============================
from docx import Document

WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
PARA_TAG = WORD_NAMESPACE + 'p'
TBL_TAG = WORD_NAMESPACE + 'tbl'
TEXT_TAG = WORD_NAMESPACE + 't'
TEXT_BOX_TAG = WORD_NAMESPACE + 'txbxContent'

def get_clean_text(xml_element):
    text_parts = []
    for node in xml_element.findall(f'.//{TEXT_TAG}'):
        if node.text:
            text_parts.append(node.text)
    return "".join(text_parts).strip()

def get_all_headers(doc):
    all_headers = []
    for section in doc.sections:
        header_data = {"floating": [], "standard": []}
        header = section.header
        xml_element = header._element

        for box in xml_element.findall(f'.//{TEXT_BOX_TAG}'):
            box_text = get_clean_text(box)
            if box_text:
                header_data["floating"].append(box_text)

        for paragraph in header.paragraphs:
            text = paragraph.text.strip()
            if text:
                header_data["standard"].append(text)

        all_headers.append(header_data)
    return all_headers

def extract_interleaved_data(file_path):
    print(f"--- PROCESSING: {file_path} ---\n")
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"Error: {e}")
        return

    section_headers = get_all_headers(doc)
    current_block_index = 0

    for element in doc.element.body:
        if element.tag == TBL_TAG:
            rows_data = []
            for row in element.findall(f'.//{WORD_NAMESPACE}tr'):
                cells = row.findall(f'.//{WORD_NAMESPACE}tc')
                row_text_list = [get_clean_text(cell) for cell in cells]
                rows_data.append(row_text_list)

            if not rows_data:
                continue

            first_row_str = " ".join(rows_data[0])

            if "Date" in first_row_str or "Session" in first_row_str:
                if current_block_index < len(section_headers):
                    h_data = section_headers[current_block_index]
                    print(f"\n[SECTION {current_block_index + 1} LAYOUT]")
                    if h_data["floating"]:
                        print(f"  • Floating Info: {', '.join(h_data['floating'])}")
                    if h_data["standard"]:
                        print(f"  • Standard Title: {h_data['standard'][0] if h_data['standard'] else ''}...")

                print(f"\n>>> BLOCK {current_block_index + 1} START <<<")
                print(f"[SECONDARY HEADER]:")
                for row in rows_data:
                    clean_row = " | ".join([c for c in row if c])
                    print(f"  {clean_row}")

                current_block_index += 1

            elif "Register Number" in first_row_str or "Sl. No" in first_row_str:
                print(f"[STUDENT LIST]:")
                print(f"  HEAD: {rows_data[0]}")
                for row in rows_data[1:]:
                    clean_row = [x for x in row if x]
                    if clean_row:
                        print(f"  DATA: {clean_row}")

        elif element.tag == PARA_TAG:
            text = get_clean_text(element)
            if not text:
                continue
            if any(k in text.upper() for k in ["INVIGILATOR", "ABSENTEES", "SUPERINTENDENT", "OFFICE USE"]):
                print(f"[FOOTER]: {text}")

# ==============================

def parse_exam_file_wrapper(file_path):
    """
    Wraps/Re-implements the logic to return structured data grouped by pages/blocks.
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        return {"error": str(e), "data": []}

    section_headers = get_all_headers(doc)
    current_block_index = 0
    
    import re

    # Structure: [ { "page_idx": 1, "room": "GBL201", "records": [...] }, ... ]
    pages = []
    current_page_records = []
    current_room = "-"
    
    # Helper to push current page
    def flush_page():
        nonlocal current_page_records, current_block_index, current_room
        if current_page_records:
            pages.append({
                "page_idx": current_block_index, 
                "room": current_room,
                "records": current_page_records
            })
            current_page_records = []

    for element in doc.element.body:
        if element.tag == TBL_TAG:
            rows_data = []
            for row in element.findall(f'.//{WORD_NAMESPACE}tr'):
                cells = row.findall(f'.//{WORD_NAMESPACE}tc')
                row_text_list = [get_clean_text(cell) for cell in cells]
                rows_data.append(row_text_list)

            if not rows_data:
                continue

            first_row_str = " ".join(rows_data[0])

            # BLOCK / PAGE BREAK LOGIC
            if "Date" in first_row_str or "Session" in first_row_str:
                flush_page()
                current_block_index += 1
                
                # Extract Room from Header
                # Image shows "Room No" "GBL301" in header area
                if current_block_index <= len(section_headers):
                     h_data = section_headers[current_block_index - 1]
                     
                     # 1. Search in floating (text boxes)
                     for text in h_data["floating"]:
                         # Check for "Room No" label AND value in same string
                         room_match = re.search(r'(?:Room|Hall)\s*(?:No\.?)?\s*[:\-]?\s*([A-Za-z0-9\-\/]+)', text, re.IGNORECASE)
                         if room_match:
                             candidate = room_match.group(1).strip()
                             if len(candidate) > 1 and candidate.lower() not in ['no', 'date', 'session', 'time']:
                                 current_room = candidate
                                 break # Found it

                         # Fallback: if text is JUST "GBL301" or similar code?
                         # Only if we saw "Room" in a previous box? (Hard to track order)
                         # But let's check if it matches a Room Code pattern directly? (Dangerous)
                         if re.match(r'^[A-Z]{2,4}[0-9]{3}$', text.strip()): # e.g. GBL301
                             current_room = text.strip()
                     
                     # 2. Search in standard paragraphs if not found in floating
                     if current_room == "-":
                         for text in h_data["standard"]:
                              room_match = re.search(r'(?:Room|Hall)\s*(?:No\.?)?\s*[:\-\.]?\s*([A-Za-z0-9\-\/]+)', text, re.IGNORECASE)
                              if room_match:
                                  candidate = room_match.group(1).strip()
                                  if len(candidate) > 1 and candidate.lower() not in ['no', 'date', 'session', 'time']:
                                      current_room = candidate
                                      break

            elif "Register Number" in first_row_str or "Sl. No" in first_row_str:
                header = [c.lower() for c in rows_data[0]]
                
                usn_idx = -1
                seat_idx = -1
                course_idx = -1
                
                for idx, col in enumerate(header):
                    if "register" in col or "usn" in col: usn_idx = idx
                    if "sl" in col and "no" in col: seat_idx = idx # Explicit: Sl. No IS Seat No
                    if "course" in col and "code" in col: course_idx = idx
                
                for row_idx, row in enumerate(rows_data[1:]): # Skip table header
                    if not any(row): continue
                    
                    usn_val = row[usn_idx] if usn_idx != -1 and usn_idx < len(row) else "UNKNOWN"
                    seat_val = row[seat_idx] if seat_idx != -1 and seat_idx < len(row) else "UNKNOWN"
                    course_val = row[course_idx] if course_idx != -1 and course_idx < len(row) else "-"
                    
                    if not usn_val or usn_val == "UNKNOWN": continue

                    record = {
                        "usn": usn_val,
                        "seat": seat_val,
                        "course": course_val,
                        "raw_data": row
                    }
                    
                    current_page_records.append(record)

    # Flush last page
    flush_page()
            
    return {"pages": pages, "total_count": sum(len(p['records']) for p in pages)}
